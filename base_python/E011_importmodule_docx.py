from docx import Document

createdDocxDocument = Document()

createdDocxDocument.add_heading('Is Python magic?')
createdDocxDocument.add_paragraph('Or am I?')

createdDocxDocument.save('createdWithPython/Is Python Magical.docx')
